# ----------------------------------------------------------------------------
# Dataiku Python Recipe  –  Recipe 2: DOWNLOAD, EXTRACT & CHUNK
# Name  : recipe_2_download_extract_chunk
# Input : plm_documents_raw        (Dataiku dataset — provided DataFrame)
# Output: plm_chunks_raw           (Dataiku dataset — one row per text chunk)
# ----------------------------------------------------------------------------
# Purpose
# -------
# For each PLM document row:
#   1. Authenticates to PLM (Windchill) via OAuth2 ROPC
#   2. Fetches PrimaryContent metadata from the PLM OData API
#   3. Downloads the file (.docx / .doc / .pdf)
#   4. Extracts text and images from the file
#   5. Uploads extracted images to S3
#   6. Chunks the extracted text (500-char, 50-char overlap)
#   7. Writes one row per chunk to the output dataset "plm_chunks_raw"
#      (with image placeholders replaced by S3 URLs)
# ----------------------------------------------------------------------------

import dataiku
import pandas as pd
import os
import json
import uuid
import base64
import logging
import shutil
import zipfile
import traceback
from datetime import datetime
from pathlib import Path
from typing import Optional
from uuid import uuid4

import boto3
import requests
import urllib3

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-8s  %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
log = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Dataiku datasets
# ---------------------------------------------------------------------------
input_dataset  = dataiku.Dataset("plm_documents_raw")
output_dataset = dataiku.Dataset("plm_chunks_raw")

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
PLM_AUTH = {
    "token_url":     os.getenv("PLM_TOKEN_URL",    "https://ilogin.infineon.com/as/token.oauth2"),
    "client_id":     os.getenv("PLM_CLIENT_ID",    "PLM ROPC"),
    "client_secret": os.getenv("PLM_CLIENT_SECRET","GjXzTdPAgqU10PmWemI44s45nbLwTV2lCVLkc3vIPrQZ5b9sStcLjpBzqR7uXnvA"),
    "username":      os.getenv("PLM_USER",         "SINqmagent"),
    "password":      os.getenv("PLM_PASSWORD",     'zMx(#LTqx5Ja+vK"$'),
    "scope":         "profile",
}

PLM_API = {
    "base_url": os.getenv("PLM_API_BASE", "https://gravitee-gw-plm.intra.infineon.com/plm/dms/oauth/servlet/odata/v7"),
    "api_key":  os.getenv("PLM_API_KEY",  "48ee3431-6f5c-4b4b-ae34-316f5ccb4b2d"),
}

S3_CONFIG = {
    "endpoint":   "https://s3muccephp.infineon.com",
    "bucket":     "quai-s3-storage",
    "access_key": "15G5PEYY3NPQTQ8OQ8I6",
    "secret_key": "CxmXDLjZMUtfWqkBmb4uHD4nxMkft3OpeilwRB6E",
}

PIPELINE_CONFIG = {
    "download_dir":      Path("/tmp/plm_downloads"),
    "image_dir":         Path("/tmp/plm_images"),
    "chunk_size":        500,
    "chunk_overlap":     50,
    "supported_formats": {".docx", ".doc", ".pdf"},
}

PIPELINE_CONFIG["download_dir"].mkdir(parents=True, exist_ok=True)
PIPELINE_CONFIG["image_dir"].mkdir(parents=True, exist_ok=True)

# ---------------------------------------------------------------------------
# S3 client
# ---------------------------------------------------------------------------
s3 = boto3.client(
    "s3",
    endpoint_url          = S3_CONFIG["endpoint"],
    aws_access_key_id     = S3_CONFIG["access_key"],
    aws_secret_access_key = S3_CONFIG["secret_key"],
    verify                = False,
)
log.info("S3 client initialised.")

# ---------------------------------------------------------------------------
# Token cache
# ---------------------------------------------------------------------------
_token_cache: dict = {}


# ===========================================================================
# Helpers
# ===========================================================================

def get_plm_access_token(force_refresh: bool = False) -> str:
    if not force_refresh and _token_cache.get("access_token"):
        return _token_cache["access_token"]
    raw = f"{PLM_AUTH['client_id']}:{PLM_AUTH['client_secret']}"
    b64 = base64.b64encode(raw.encode()).decode()
    headers = {
        "Authorization": f"Basic {b64}",
        "Content-Type":  "application/x-www-form-urlencoded",
    }
    payload = {
        "grant_type": "password",
        "username":   PLM_AUTH["username"],
        "password":   PLM_AUTH["password"],
        "scope":      PLM_AUTH["scope"],
    }
    resp = requests.post(PLM_AUTH["token_url"], headers=headers, data=payload,
                         timeout=30, verify=False)
    resp.raise_for_status()
    token = resp.json()["access_token"]
    _token_cache["access_token"] = token
    log.info("PLM access token acquired.")
    return token


def plm_headers() -> dict:
    try:
        token = get_plm_access_token()
        return {
            "Content-Type":       "application/json",
            "X-Gravitee-Api-Key": PLM_API["api_key"],
            "Authorization":      f"Bearer {token}",
            "accept":             "*/*",
        }
    except Exception as exc:
        log.error("Failed to build PLM headers: %s", exc)
        return {"Content-Type": "application/json", "X-Gravitee-Api-Key": PLM_API["api_key"]}


def get_primary_content_meta(wt_document_id: str) -> Optional[dict]:
    try:
        if not wt_document_id.startswith("OR:"):
            wt_document_id = f"OR:wt.doc.WTDocument:{wt_document_id}"
        url = f"{PLM_API['base_url']}/DocMgmt/Documents('{wt_document_id}')/PrimaryContent"
        resp = requests.get(url, headers=plm_headers(), timeout=30, verify=False)
        if resp.status_code == 401:
            log.warning("PLM token expired – refreshing …")
            get_plm_access_token(force_refresh=True)
            resp = requests.get(url, headers=plm_headers(), timeout=30, verify=False)
        resp.raise_for_status()
        return resp.json()
    except Exception as exc:
        log.error("Error fetching PrimaryContent for %s: %s", wt_document_id, exc)
        return None


def download_document(download_url: str, dest_path: Path) -> bool:
    try:
        resp = requests.get(download_url, headers=plm_headers(),
                            stream=True, timeout=120, verify=False)
        resp.raise_for_status()
        with dest_path.open("wb") as fh:
            for chunk in resp.iter_content(chunk_size=65536):
                fh.write(chunk)
        return True
    except Exception as exc:
        log.error("Download failed for %s: %s", download_url, exc)
        return False


def extract_text_docx(file_path: Path) -> tuple:
    from docx import Document as DocxDocument
    import xml.etree.ElementTree as ET

    image_list = []
    image_output_dir = PIPELINE_CONFIG["image_dir"]

    try:
        doc = DocxDocument(str(file_path))
        image_map = {}

        with zipfile.ZipFile(str(file_path), "r") as docx_zip:
            rels_path = "word/_rels/document.xml.rels"
            if rels_path in docx_zip.namelist():
                rels_xml  = docx_zip.read(rels_path)
                rels_tree = ET.fromstring(rels_xml)
                ns         = "http://schemas.openxmlformats.org/package/2006/relationships"
                image_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
                for rel in rels_tree.findall(f"{{{ns}}}Relationship"):
                    if rel.attrib.get("Type") == image_type:
                        r_id        = rel.attrib["Id"]
                        target      = rel.attrib["Target"]
                        full_target = f"word/{target}"
                        if full_target in docx_zip.namelist():
                            suffix      = Path(target).suffix
                            unique_name = f"{uuid4().hex}{suffix}"
                            save_path   = image_output_dir / unique_name
                            with docx_zip.open(full_target) as img_file:
                                with open(save_path, "wb") as out_file:
                                    shutil.copyfileobj(img_file, out_file)
                            image_map[r_id] = unique_name
                            image_list.append(unique_name)

        parts = []
        for para in doc.paragraphs:
            para_parts = []
            for run in para.runs:
                if run.text:
                    para_parts.append(run.text)
                blips = run._r.findall(
                    ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
                )
                for blip in blips:
                    r_embed = blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    )
                    if r_embed and r_embed in image_map:
                        para_parts.append(f"[IMAGE:{image_map[r_embed]}]")
            if para_parts:
                parts.append("".join(para_parts))

        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    parts.append(" | ".join(row_texts))

        return "\n".join(parts), image_list

    except Exception as exc:
        log.error("Error extracting docx %s: %s", file_path, exc)
        return "", image_list


def extract_text_pdf(file_path: Path) -> tuple:
    import fitz  # PyMuPDF

    image_list = []
    image_output_dir = PIPELINE_CONFIG["image_dir"]
    parts = []

    try:
        doc = fitz.open(str(file_path))
        for page in doc:
            parts.append(page.get_text())
            for img_index, img in enumerate(page.get_images(full=True)):
                base_image  = doc.extract_image(img[0])
                unique_name = f"{uuid4().hex}.{base_image['ext']}"
                save_path   = image_output_dir / unique_name
                save_path.write_bytes(base_image["image"])
                image_list.append(unique_name)
        doc.close()
        return "\n".join(parts), image_list
    except Exception as exc:
        log.error("Error extracting PDF %s: %s", file_path, exc)
        return "", image_list


def extract_text(file_path: Path) -> tuple:
    ext = file_path.suffix.lower()
    if ext in (".docx", ".doc"):
        return extract_text_docx(file_path)
    elif ext == ".pdf":
        return extract_text_pdf(file_path)
    return "", []


def upload_images_to_s3(doc_id: str, image_list: list) -> tuple:
    uploaded_images = {}
    failed_images   = []
    image_output_dir = PIPELINE_CONFIG["image_dir"]

    magic_bytes_map = {
        b'\xff\xd8\xff':      ('image/jpeg', '.jpg'),
        b'\x89PNG\r\n\x1a\n': ('image/png',  '.png'),
        b'GIF87a':             ('image/gif',  '.gif'),
        b'GIF89a':             ('image/gif',  '.gif'),
        b'BM':                 ('image/bmp',  '.bmp'),
        b'II\x2a\x00':        ('image/tiff', '.tiff'),
        b'MM\x00\x2a':        ('image/tiff', '.tiff'),
        b'RIFF':               ('image/webp', '.webp'),
    }

    for idx, imagepath in enumerate(image_list):
        try:
            image_path  = image_output_dir / imagepath
            final_bytes = image_path.read_bytes()
            if not final_bytes:
                raise ValueError("Empty image file.")

            ext          = image_path.suffix.lower() or ".png"
            content_type = f"image/{ext.lstrip('.')}"
            for magic, (mime, extension) in magic_bytes_map.items():
                if final_bytes.startswith(magic):
                    content_type = mime
                    ext          = extension
                    break

            image_id = f"{idx}_{image_path.stem}{ext}"
            s3_key   = f"production_document/images/{doc_id}/{image_id}"

            s3.put_object(
                Bucket=S3_CONFIG["bucket"],
                Key=s3_key,
                Body=final_bytes,
                ContentType=content_type,
            )
            s3_url = f"{S3_CONFIG['endpoint']}/{S3_CONFIG['bucket']}/{s3_key}"
            uploaded_images[imagepath] = s3_url

        except Exception as exc:
            log.error("Image upload failed (idx=%d, doc=%s): %s", idx, doc_id, exc)
            failed_images.append({"index": idx, "reason": str(exc)})

    return uploaded_images, failed_images


def chunk_text(text: str, chunk_size: int, overlap: int) -> list:
    chunks, start = [], 0
    while start < len(text):
        chunks.append(text[start:start + chunk_size])
        start += chunk_size - overlap
    return chunks


# ===========================================================================
# Per-document processing
# ===========================================================================

def process_document(doc_row: dict, run_id: str) -> list:
    """
    Download, extract, upload images, chunk.
    Returns a list of chunk dicts (one per chunk) — empty list if skipped/failed.
    """
    doc_id   = str(doc_row.get("id", ""))
    doc_num  = doc_row.get("doc_number", "")
    filename = doc_row.get("filename", "") or ""

    log.info("Processing  id=%-15s  number=%-15s  file=%s", doc_id, doc_num, filename)

    try:
        content_meta = get_primary_content_meta(doc_id)
        if not content_meta:
            log.warning("Skipping %s – no PrimaryContent metadata.", doc_id)
            return []

        file_name    = content_meta.get("FileName", filename)
        download_url = (content_meta.get("Content") or {}).get("URL")
        if not download_url:
            log.warning("Skipping %s – no download URL.", doc_id)
            return []

        file_ext = Path(file_name).suffix.lower()
        if file_ext not in PIPELINE_CONFIG["supported_formats"]:
            log.info("Skipping %s – unsupported format '%s'.", doc_id, file_ext)
            return []

        # Download file
        local_path = PIPELINE_CONFIG["download_dir"] / f"{doc_id}{file_ext}"
        if not download_document(download_url, local_path):
            return []

        # Extract text + images
        full_text, image_list = extract_text(local_path)
        if not full_text.strip():
            log.warning("No text extracted from %s.", local_path)

        # Upload images to S3
        uploaded_map, _ = upload_images_to_s3(doc_id, image_list)

        # Chunk text and replace image placeholders with S3 URLs
        chunks     = chunk_text(full_text, PIPELINE_CONFIG["chunk_size"], PIPELINE_CONFIG["chunk_overlap"])
        chunk_rows = []

        for chunk_seq, text in enumerate(chunks):
            for placeholder, s3_url in uploaded_map.items():
                text = text.replace(placeholder, s3_url)

            chunk_rows.append({
                "run_id":       run_id,
                "doc_id":       doc_id,
                "doc_number":   doc_num,
                "filename":     file_name,
                "filepath":     download_url,
                "filetype":     file_ext.lstrip("."),
                "chunk_seq":    chunk_seq,
                "chunk_text":   text,
                "images_count": len(uploaded_map),
                "chunked_at":   datetime.utcnow().isoformat(),
            })

        # Clean up local download
        try:
            local_path.unlink(missing_ok=True)
        except Exception:
            pass

        log.info("doc_id=%s → %d chunks, %d images uploaded.", doc_id, len(chunk_rows), len(uploaded_map))
        return chunk_rows

    except Exception as exc:
        log.error("Unhandled error processing %s: %s\n%s", doc_id, exc, traceback.format_exc())
        return []


# ===========================================================================
# Main
# ===========================================================================

def main():
    run_id = uuid.uuid4().hex[:12]
    log.info("=" * 60)
    log.info("Recipe 2: Download, Extract & Chunk  run_id=%s", run_id)
    log.info("=" * 60)

    df = input_dataset.get_dataframe()
    log.info("Read %d rows from 'plm_documents_raw'.", len(df))

    if df.empty:
        log.warning("Input dataset is empty – nothing to process.")
        output_dataset.write_with_schema(pd.DataFrame())
        return

    # Authenticate to PLM upfront
    try:
        get_plm_access_token()
    except Exception as exc:
        log.error("PLM authentication failed – aborting: %s", exc)
        raise

    all_chunks = []
    total = len(df)

    for idx, row in df.iterrows():
        try:
            chunk_rows = process_document(row.to_dict(), run_id)
            all_chunks.extend(chunk_rows)
        except Exception as exc:
            log.error("Fatal error on row %d: %s", idx, exc)

        if (idx + 1) % 50 == 0:
            log.info("Progress: %d / %d documents …", idx + 1, total)

    chunks_df = pd.DataFrame(all_chunks)
    output_dataset.write_with_schema(chunks_df)

    log.info("=" * 60)
    log.info("Recipe 2 complete.  Documents=%d  Total chunks=%d", total, len(chunks_df))
    log.info("=" * 60)


main()
